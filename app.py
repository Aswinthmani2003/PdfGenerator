from docx import Document
from datetime import datetime
import os
import streamlit as st
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
import fitz  # PyMuPDF
import tempfile

port = int(os.environ.get("PORT", 8501))
# Function to generate a unique reference number
def generate_reference_number(company_name="BKR"):
    current_month = datetime.now().strftime("%m")
    current_year = datetime.now().strftime("%Y")
    serial_number = datetime.now().strftime("%d%H%M%S")  # Unique serial
    return f"{company_name}{current_month}-{current_year}-CR{serial_number}"

# Function to replace placeholders
def replace_placeholders(doc, placeholders):
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in placeholders.items():
                        if key in para.text:
                            para.text = para.text.replace(key, value)
    return doc

# Streamlit UI
st.title("Document Generator")
template_option = st.selectbox("Select Template", ["Service Agreement", "Invoice", "VAT Registration"])

if template_option == "Service Agreement":
    agreement_date = st.date_input("Date of Agreement", datetime.today())
    client_name = st.text_input("Client Name")
    signatory_name = st.text_input("Signatory Name")
    passport_number = st.text_input("Passport Number")
    reference_number = generate_reference_number()
    
    # Input costs
    costs = {
        "{COMPANY_FORMATION_COST}": st.number_input("Company Formation Cost", min_value=0.0, step=0.01),
        "{OFFICE_RENTAL_COST}": st.number_input("Office Rental Cost", min_value=0.0, step=0.01),
        "{VISA_COST}": st.number_input("Visa Cost", min_value=0.0, step=0.01),
        "{ADMIN_CHARGES}": st.number_input("Admin Charges", min_value=0.0, step=0.01),
        "{POWER_OF_ATTORNEY_COST}": st.number_input("Power of Attorney Cost", min_value=0.0, step=0.01),
        "{BUSINESS_GUIDANCE_COST}": st.number_input("Business Guidance Cost", min_value=0.0, step=0.01),
    }
    total_cost = sum(costs.values())
    
    placeholders = {
        "{REFERENCE_NUMBER}": reference_number,
        "{CLIENT_NAME}": client_name,
        "{SIGNATORY_NAME}": signatory_name,
        "{SIGNATORY_PASSPORT_NUMBER}": passport_number,
        "{TOTAL_COST}": f"{total_cost:.2f}"
    }
    placeholders.update(costs)
    
    if st.button("Generate Service Agreement"):
        try:
            doc_path = "SAMPLE Service Agreement -Company formation -Bahrain - Filled (1).docx"
            doc = Document(doc_path)
            doc = replace_placeholders(doc, placeholders)
            output_path = "Service_Agreement_Generated.docx"
            doc.save(output_path)
            st.success("Service Agreement generated successfully!")
            with open(output_path, "rb") as file:
                st.download_button("Download Service Agreement", file, file_name=output_path)
        except Exception as e:
            st.error(f"Error: {e}")

elif template_option == "Invoice":
    invoice_date = st.date_input("Invoice Date", datetime.today())
    client_name = st.text_input("Client Name")
    reference_number = st.text_input("Service Agreement Reference Number")
    service = st.text_input("Service Description")
    cost = st.number_input("Cost (in BHD)", min_value=0.0, step=0.01)
    total_amount = f"{cost:.2f}"
    invoice_number = generate_reference_number("INV")
    
    placeholders = {
        "{INVOICE_DATE}": invoice_date.strftime("%d-%m-%Y"),
        "{INVOICE_NUMBER}": invoice_number,
        "{CLIENT_NAME}": client_name,
        "{REFERENCE_NUMBER}": reference_number,
        "{SERVICE_DESCRIPTION}": service,
        "{TOTAL_AMOUNT}": total_amount,
    }
    
    if st.button("Generate Invoice"):
        try:
            doc_path = "SAMPLE -Invoice BKR2024CF158 - first payment.docx"
            doc = Document(doc_path)
            doc = replace_placeholders(doc, placeholders)
            output_path = "Invoice_Generated.docx"
            doc.save(output_path)
            st.success("Invoice generated successfully!")
            with open(output_path, "rb") as file:
                st.download_button("Download Invoice", file, file_name=output_path)
        except Exception as e:
            st.error(f"Error: {e}")
