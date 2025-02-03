import streamlit as st
from docx import Document
from datetime import datetime
import os
import platform
import subprocess
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

# Function to generate a unique reference number
def generate_reference_number(company_name="BKR"):
    current_month = datetime.now().strftime("%m")
    current_year = datetime.now().strftime("%Y")
    serial_number = datetime.now().strftime("%d%H%M%S")  # Unique serial
    return f"{company_name}{current_month}-{current_year}-CR{serial_number}"

# Function to replace placeholders in Word document
def replace_placeholders(doc, placeholders):
    """Replace placeholders in a Word document, including paragraphs and tables."""
    
    def replace_in_paragraph(paragraph, key, value):
        """Replace placeholders in a single paragraph, handling split runs."""
        full_text = "".join(run.text for run in paragraph.runs)
        if key in full_text:
            full_text = full_text.replace(key, value)
            for run in paragraph.runs:
                run.text = ""  # Clear all runs
            paragraph.runs[0].text = full_text  # Add the replaced text back

    def replace_in_cell(cell, placeholders):
        """Replace placeholders inside a table cell."""
        for para in cell.paragraphs:
            for key, value in placeholders.items():
                replace_in_paragraph(para, key, value)

    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_in_paragraph(para, key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, placeholders)

    return doc

# Function to convert DOCX to PDF
def convert_docx_to_pdf(docx_path, pdf_path):
    """Converts DOCX to PDF while retaining formatting."""
    if platform.system() == "Windows":
        import comtypes.client
        word = comtypes.client.CreateObject("Word.Application")
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()
    elif platform.system() == "Linux":
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", docx_path])
    return pdf_path

# Streamlit UI
st.title("Service Agreement Generator")

# Input fields
reference_number = generate_reference_number()
client_name = st.text_input("Client Name")
signatory_name = st.text_input("Signatory Name")
passport_number = st.text_input("Passport Number")
ISIC_CODE_1 = st.text_input("Business Activity Code 1")
BUSINESS_ACTIVITY_1 = st.text_input("Business Activity Name 1")
BUSINESS_DESCRIPTION_1 = st.text_area("Business Description 1")
ISIC_CODE_2 = st.text_input("Business Activity Code 2")
BUSINESS_ACTIVITY_2 = st.text_input("Business Activity Name 2")
BUSINESS_DESCRIPTION_2 = st.text_area("Business Description 2")

# Input costs
costs = {
    "{COMPANY_FORMATION_COST}": st.number_input("Company Formation Cost", min_value=0.0, step=0.01),
    "{OFFICE_RENTAL_COST}": st.number_input("Office Rental Cost", min_value=0.0, step=0.01),
    "{VISA_COST}": st.number_input("Visa Cost", min_value=0.0, step=0.01),
    "{ADMIN_CHARGES}": st.number_input("Administrative Charges", min_value=0.0, step=0.01),
    "{POWER_OF_ATTORNEY_COST}": st.number_input("Power of Attorney Cost", min_value=0.0, step=0.01),
    "{PRO_SERVICES_COST}": st.number_input("PRO Services Cost", min_value=0.0, step=0.01),
    "{LABOUR_REGISTRATION_COST}": st.number_input("Labour Registration Cost", min_value=0.0, step=0.01),
    "{SOCIAL_INSURANCE_COST}": st.number_input("Social Insurance Registration Cost", min_value=0.0, step=0.01),
    "{BUSINESS_GUIDANCE_COST}": st.number_input("Business Guidance Cost", min_value=0.0, step=0.01),
}
total_cost = sum(costs.values())

# Placeholders
placeholders = {
    "{REFERENCE_NUMBER}": reference_number,
    "{CLIENT_NAME}": client_name,
    "{SIGNATORY_NAME}": signatory_name,
    "{SIGNATORY_PASSPORT_NUMBER}": passport_number,
    "{ISIC_CODE_1}": ISIC_CODE_1,
    "{BUSINESS_ACTIVITY_1}": BUSINESS_ACTIVITY_1,
    "{BUSINESS_DESCRIPTION_1}": BUSINESS_DESCRIPTION_1,
    "{ISIC_CODE_2}": ISIC_CODE_2,
    "{BUSINESS_ACTIVITY_2}": BUSINESS_ACTIVITY_2,
    "{BUSINESS_DESCRIPTION_2}": BUSINESS_DESCRIPTION_2,
    "{TOTAL_COST}": f"{total_cost:.2f}",
}
placeholders.update({k: f"{v:.2f}" for k, v in costs.items()})  # Add formatted costs to placeholders

# Generate the document
if st.button("Generate Service Agreement"):
    try:
        doc_path = "SAMPLE Service Agreement -Company formation -Bahrain - Filled (1).docx"
        doc = Document(doc_path)
        doc = replace_placeholders(doc, placeholders)
        
        # Save as DOCX
        docx_output_path = "Service_Agreement_Generated.docx"
        doc.save(docx_output_path)

        # Convert DOCX to PDF
        pdf_output_path = "Service_Agreement_Generated.pdf"
        convert_docx_to_pdf(docx_output_path, pdf_output_path)

        st.success("Service Agreement generated successfully!")

        # Separate buttons for Word and PDF download
        col1, col2 = st.columns(2)

        with col1:
            with open(docx_output_path, "rb") as file:
                st.download_button("Download as Word", file, file_name="Service_Agreement.docx")

        with col2:
            with open(pdf_output_path, "rb") as file:
                st.download_button("Download as PDF", file, file_name="Service_Agreement.pdf")

    except Exception as e:
        st.error(f"Error: {e}")
